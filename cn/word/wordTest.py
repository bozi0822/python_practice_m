from datetime import datetime
from operator import contains
from icecream import ic
from docx import Document
from docx.shared import Inches

path = r"文档.docx"
document = Document(path)


# def time_format():
#     return f'{datetime.now()}|> '
#
#
# ic.configureOutput(prefix=time_format)

# 输出段落编号及段落内容
count = 0
for i in range(len(document.paragraphs)):
    content = document.paragraphs[i].text
    if contains(content, "word文档"):
        count += 1
    ic("第" + str(i) + "段的内容是：" + content)

ic(count)
